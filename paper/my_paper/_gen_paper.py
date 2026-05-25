"""
Generate thesis docx for whr-im IM system.
Format: A4, margins per spec, body 宋体 小四 (12pt), 1.5 line spacing.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

ROOT = '/Users/zyb/bishe/whr-im/paper/my_paper'
PIC = os.path.join(ROOT, 'picture')
DIA = os.path.join(PIC, 'diagrams')

# Paper title (used in header and front matter)
PAPER_TITLE = '基于Gin和WebSocket的IM即时通讯系统'


def set_run_font(run, cn='宋体', en='Times New Roman', size=12, bold=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)


def add_para(doc, text, size=12, bold=False, align=None, indent_first=True,
             line_spacing=1.5, space_before=0, space_after=0, cn='宋体', en='Times New Roman'):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Pt(size * 2)  # 2 chars indent
    run = p.add_run(text)
    set_run_font(run, cn=cn, en=en, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level <= 2 else 6)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, cn='黑体', en='Times New Roman', size=sizes.get(level, 12), bold=True)
    return p


def add_centered(doc, text, size=12, bold=False, space_before=0, space_after=0, cn='宋体'):
    return add_para(doc, text, size=size, bold=bold,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    indent_first=False, space_before=space_before,
                    space_after=space_after, cn=cn)


def add_image(doc, path, caption, width_inch=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inch))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, cn='宋体', en='Times New Roman', size=10.5, bold=True)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    r = p.add_run(code)
    set_run_font(r, cn='宋体', en='Courier New', size=10)


# ----------------- Section / Header / Footer / Page Number -----------------
def _make_section_props(section, page_num_fmt='decimal', start=1,
                        title_only_first_section=False):
    """Set page-number format on a section's sectPr.

    page_num_fmt: 'decimal' (1,2,3) | 'upperRoman' (I,II,III) | 'lowerRoman'
    """
    sectPr = section._sectPr
    # Remove existing pgNumType
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:fmt'), page_num_fmt)
    el.set(qn('w:start'), str(start))
    sectPr.append(el)


def _set_header_text(section, text, size=9):
    """Set header paragraph text centered. size=9pt (小五).
    Header中文体宋体小五号字居中。"""
    h = section.header
    h.is_linked_to_previous = False
    # Clear existing
    for p in list(h.paragraphs):
        if p.text:
            for r in list(p.runs):
                r.text = ''
        # do not remove the paragraph itself
    # Use the first paragraph
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # remove all existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(text)
    set_run_font(r, cn='宋体', en='Times New Roman', size=size)


def _clear_header(section):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def _set_footer_page_number(section, size=9, en='Times New Roman'):
    """Insert PAGE field centered into footer."""
    f = section.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    # Insert: <w:fldChar fldCharType="begin"/><w:instrText>PAGE</w:instrText>
    #         <w:fldChar fldCharType="end"/>
    run = p.add_run()
    set_run_font(run, cn='宋体', en=en, size=size)
    r_el = run._element

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE   \\* MERGEFORMAT '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r_el.append(fld_begin)
    r_el.append(instr)
    r_el.append(fld_end)


def _clear_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def add_new_section(doc, page_num_fmt='decimal', start=1, header_text=None,
                    show_page_number=True):
    """Add a new section (next-page break) and configure header/footer.

    Returns the new section.
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    # Page setup (inherit but be explicit)
    new_section.page_height = Cm(29.7)
    new_section.page_width = Cm(21.0)
    new_section.top_margin = Cm(2.5)
    new_section.bottom_margin = Cm(2.5)
    new_section.left_margin = Cm(3.0)
    new_section.right_margin = Cm(2.0)
    new_section.header_distance = Cm(1.5)
    new_section.footer_distance = Cm(1.75)

    _make_section_props(new_section, page_num_fmt=page_num_fmt, start=start)
    if header_text:
        _set_header_text(new_section, header_text)
    else:
        _clear_header(new_section)
    if show_page_number:
        _set_footer_page_number(new_section)
    else:
        _clear_footer(new_section)
    return new_section


def setup_doc():
    doc = Document()
    # Default Section 0: cover/任务书 — no header, no footer.
    for s in doc.sections:
        s.page_height = Cm(29.7)
        s.page_width = Cm(21.0)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.0)
        s.header_distance = Cm(1.5)
        s.footer_distance = Cm(1.75)
        _clear_header(s)
        _clear_footer(s)

    # Default Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return doc


def add_page_break(doc):
    doc.add_page_break()


if __name__ == '__main__':
    doc = setup_doc()
    out = os.path.join(ROOT, '我的毕业论文.docx')
    doc.save(out)
    print('Skeleton saved:', out)

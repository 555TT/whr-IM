"""Cover + 任务书 + 摘要 + ABSTRACT + 目录"""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from _gen_paper import (
    add_para,
    add_heading,
    add_centered,
    add_page_break,
    add_new_section,
    add_blank_line,
    PAPER_TITLE,
    set_run_font,
)


def _cover_line(doc, label, value='', underline_value=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(32)
    p.paragraph_format.right_indent = Pt(10)
    r1 = p.add_run(f'{label}    ')
    set_run_font(r1, cn='宋体', size=16)
    r2 = p.add_run(value)
    set_run_font(r2, cn='黑体' if underline_value else '宋体', size=16)
    if underline_value:
        r2.font.underline = True
    # 补一段空白，使视觉效果更像模板
    r3 = p.add_run(' ' * 46)
    set_run_font(r3, cn='宋体', size=16)


def write_front(doc):
    # ========== 封面 ==========
    add_blank_line(doc)
    add_blank_line(doc)
    add_centered(doc, '郑州轻工业大学', size=22, bold=True, cn='宋体')
    add_centered(doc, '本科毕业设计(论文)', size=36, bold=False, cn='宋体')
    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)

    _cover_line(doc, '题    目', PAPER_TITLE, underline_value=True)
    _cover_line(doc, '学生姓名', '王浩然')
    _cover_line(doc, '专业班级', '软件工程22-0X')
    _cover_line(doc, '学    号', '5421220XXXX')
    _cover_line(doc, '学    院', '软件学院')
    _cover_line(doc, '指导教师（职称）', '***（***）')
    _cover_line(doc, '完成时间', '2026年5月18日')

    add_page_break(doc)

    # ========== 任务书 ==========
    add_blank_line(doc)
    add_centered(doc, '郑州轻工业大学', size=16, bold=False, cn='宋体')
    add_centered(doc, '毕业设计（论文）任务书', size=15, bold=True, cn='宋体')
    add_blank_line(doc)
    add_para(doc, f'题目：{PAPER_TITLE}', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '专业：软件工程22-0X    学号：5421220XXXX    姓名：王浩然', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '主要内容：', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '本课题围绕一套前后端分离的即时通讯系统展开。前端使用 Vue 3、Vite 与 TypeScript，'
                  '后端采用 Go、Gin、GORM 和 MySQL，先把注册登录、个人资料维护、好友申请与管理、'
                  '基于 WebSocket 的实时聊天以及历史消息查询这些基础功能跑通。真正要解决的问题，'
                  '是怎样把 RSA-OAEP 端到端消息加密机制落到浏览器侧：消息先在用户本地完成加密，再上传到服务端，'
                  '服务端只负责保存和转发密文，从而尽量压缩消息明文暴露面。',
             indent_first=True, line_spacing=1.5)
    add_para(doc, '基本要求：', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '（1）熟悉前后端分离的开发模式，能够独立完成前端工程与后端工程的搭建与联调；', line_spacing=1.5)
    add_para(doc, '（2）熟练掌握 WebSocket 协议与 JWT 鉴权机制，实现安全可靠的实时消息推送；', line_spacing=1.5)
    add_para(doc, '（3）熟悉 Web Crypto API 与 RSA-OAEP 算法原理，完成消息加密与解密功能；', line_spacing=1.5)
    add_para(doc, '（4）完成系统的需求分析、概要设计、详细设计、编码实现与测试，撰写完整的毕业论文。', line_spacing=1.5)
    add_para(doc, '思政要求：', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '在课题完成过程中，培养严谨求实的科学态度、自主学习与创新精神；通过研究消息加密技术，增强网络空间安全意识与社会责任感，树立正确的网络伦理观念。', line_spacing=1.5)
    add_para(doc, '主要参考资料：', size=14, bold=True, indent_first=False, line_spacing=1.0)
    add_para(doc, '[1] Tanenbaum A S, Wetherall D J. Computer Networks[M]. 5th ed. Boston: Pearson, 2011.', indent_first=False, hanging=True, line_spacing=1.5)
    add_para(doc, '[2] Donovan A A A, Kernighan B W. The Go Programming Language[M]. Boston: Addison-Wesley, 2015.', indent_first=False, hanging=True, line_spacing=1.5)
    add_para(doc, '[3] You E. Vue.js 3 Documentation[EB/OL]. (2024-01-01)[2026-03-01]. https://vuejs.org/.', indent_first=False, hanging=True, line_spacing=1.5)
    add_blank_line(doc)
    add_para(doc, '完  成  期  限：  2026年5月18日', size=14, bold=True, indent_first=False)
    add_para(doc, '指导教师签名：', size=14, bold=True, indent_first=False)
    add_para(doc, '专业负责人签名：', size=14, bold=True, indent_first=False)
    add_centered(doc, '2026年1月6日', size=14, bold=True, cn='宋体')

    # ========== 节2：摘要/ABSTRACT/目录 — 页眉=题目，页码=罗马数字 ==========
    add_new_section(doc, page_num_fmt='upperRoman', start=1,
                    header_text=PAPER_TITLE, show_page_number=True)

    # ========== 中文摘要 ==========
    add_centered(doc, '摘  要', size=15, bold=False, cn='宋体')
    add_blank_line(doc)
    add_para(doc, '聊天软件大多会用 TLS 把消息在传输路上保护好，但一进服务端，几乎都还是明文。数据库一旦泄露，'
                  '泄出去的就是完整的聊天记录——这件事在过去几年的安全事故里反复出现过。本文从这个具体缺口入手，'
                  '搭了一套前后端分离的 IM 系统，框架选了 Gin 加 WebSocket。注册登录、好友、实时聊天、历史消息这些功能'
                  '都按常规做法跑通，重点反而是把端到端加密真正塞进业务链路里——而不是当成额外的"加密会话"挂在一旁。')
    add_para(doc, '前端用的是 Vue 3，加上 Vite、TypeScript 和 Pinia；后端走 Go 1.22 配合 Gin、GORM、MySQL，'
                  '实时下行交给 Gorilla WebSocket 维护长连接。最花精力的还是加密这一块。算法选了 RSA-OAEP，配合浏览器自带的'
                  ' Web Crypto API——密钥对在浏览器里生成，私钥从不出本机，传到服务端的只有公钥。每发出一条消息，'
                  '前端会在本地做两次加密：一次用对方公钥，一次用自己的公钥。服务端拿到的只是两份密文，做的事也就剩校验、'
                  '入库和转发。聊天历史在两边设备上都还能读，但服务端从头到尾看不到明文。')
    add_para(doc, '功能上准备了 12 项用例，覆盖了登录、好友、聊天、历史消息和加密路径，整套跑下来 12 项全过；同时去抓包、'
                  '去数据库里翻，也没看到任何漏出来的明文。换个说法：端到端加密最核心的那一条承诺——服务端没法读它转发的消息——'
                  '在这份实现里是成立的，双方又都能本地回看历史。还有几件事文中没掩盖：公钥真实性还没做校验、暂时没有前向保密、'
                  '会话密钥协商也没引入、多设备同步同样还没覆盖。这几条就是这套系统要继续往工业级 E2EE 靠近时绕不开的下一步。')
    add_blank_line(doc)
    add_para(doc, '关键词：即时通讯；前后端分离；WebSocket；端到端加密；RSA-OAEP',
             indent_first=False, bold=False, line_spacing=1.5)

    add_blank_line(doc)
    add_blank_line(doc)
    add_blank_line(doc)

    # ========== ABSTRACT ==========
    add_centered(doc, 'ABSTRACT', size=14, bold=True, cn='Times New Roman')
    add_blank_line(doc)
    add_blank_line(doc)
    add_para(doc, 'Most IM systems today rely on TLS to protect messages in transit, yet leave them in plaintext once '
                  'they hit the database. A single server breach is enough to expose every conversation that ever passed '
                  'through. This thesis takes that gap as its starting point. It builds a front-end / back-end separated '
                  'IM system on top of Gin and WebSocket, and pushes the security boundary all the way into the browser: '
                  'messages are encrypted before they leave the sender\'s machine, so what the server sees and stores is '
                  'only ciphertext. Authentication, friend management, real-time chat and message history are still '
                  'there, but the real focus is making end-to-end encryption work inside an ordinary web stack rather '
                  'than as a separate "secure mode".', en='Times New Roman', cn='Times New Roman')
    add_para(doc, 'On the front-end side the stack is Vue 3 with Vite, TypeScript and Pinia; on the back-end it is Go '
                  '1.22 plus Gin, GORM and MySQL, with Gorilla WebSocket handling the real-time push. The encryption '
                  'piece is the part that took the most thought. RSA-OAEP was chosen together with the browser-native '
                  'Web Crypto API — key pairs are generated in the browser, the private key never leaves it, and only '
                  'the public key gets uploaded. When a message is sent, the plaintext is encrypted twice in the '
                  'browser, once under the sender\'s own public key and once under the receiver\'s. The server only ever '
                  'touches ciphertext, but because both sides keep a copy they can decrypt, history remains readable on '
                  'either device.', en='Times New Roman', cn='Times New Roman')
    add_para(doc, 'Twelve functional cases were run end to end, covering login, friends, chat, history and the '
                  'encryption path itself. All twelve passed, and neither packet captures nor a direct look at the '
                  'database turned up any leftover plaintext. Put differently: the most important promise of end-to-end '
                  'encryption — that the server cannot read the messages it forwards — does hold in this build, while '
                  'both parties can still re-read their history locally. A few problems are left open and stated '
                  'plainly: public key authenticity is not yet verified, there is no forward secrecy, no session key '
                  'negotiation, and no multi-device sync. They are the obvious next steps if anyone wants to push this '
                  'design closer to an industrial-grade E2EE system.', en='Times New Roman', cn='Times New Roman')
    add_blank_line(doc)
    add_para(doc, 'KEY WORDS: Instant Messaging; Front-end and Back-end Separation; WebSocket; End-to-End Encryption; RSA-OAEP',
             indent_first=False, bold=True, line_spacing=1.5, en='Times New Roman', cn='Times New Roman')

    # ========== 目录 ==========
    add_new_section(doc, page_num_fmt='upperRoman', start=1,
                    header_text=PAPER_TITLE, show_page_number=True)
    add_centered(doc, '目  录', size=16, bold=False, cn='宋体')
    add_blank_line(doc)
    toc_items = [
        ('摘  要', 'I'),
        ('ABSTRACT', 'II'),
        ('1  绪论', '1'),
        ('    1.1  研究背景与意义', '1'),
        ('    1.2  国内外研究现状', '2'),
        ('        1.2.1  国外研究现状', '2'),
        ('        1.2.2  国内研究现状', '3'),
        ('    1.3  论文主要工作', '3'),
        ('    1.4  论文组织结构', '4'),
        ('2  相关技术介绍', '5'),
        ('    2.1  前后端分离架构', '5'),
        ('    2.2  Vue 3 与前端技术栈', '5'),
        ('    2.3  Go 语言与 Gin 框架', '6'),
        ('    2.4  WebSocket 协议', '6'),
        ('    2.5  JWT 身份认证', '7'),
        ('    2.6  RSA-OAEP 公钥加密与 Web Crypto API', '7'),
        ('    2.7  MySQL 数据库与 GORM', '8'),
        ('3  系统需求分析与总体设计', '9'),
        ('    3.1  需求分析', '9'),
        ('        3.1.1  功能性需求', '9'),
        ('        3.1.2  非功能性需求', '10'),
        ('        3.1.3  用例分析', '10'),
        ('    3.2  系统总体架构', '11'),
        ('    3.3  系统功能模块设计', '12'),
        ('    3.4  数据库设计', '13'),
        ('4  关键技术实现', '15'),
        ('    4.1  用户认证与 JWT 实现', '15'),
        ('    4.2  好友关系处理', '16'),
        ('    4.3  基于 WebSocket 的实时聊天实现', '17'),
        ('        4.3.1  WebSocket 连接建立与鉴权', '17'),
        ('        4.3.2  连接管理与消息广播', '18'),
        ('        4.3.3  消息收发完整时序', '19'),
        ('    4.4  端到端消息加密机制', '20'),
        ('        4.4.1  设计目标与总体思路', '20'),
        ('        4.4.2  RSA-OAEP 密钥对生成与管理', '21'),
        ('        4.4.3  公钥上传与好友公钥获取', '22'),
        ('        4.4.4  双重密文加密与服务端密文转发', '22'),
        ('        4.4.5  接收端解密与历史消息展示', '23'),
        ('        4.4.6  安全性分析', '24'),
        ('5  系统实现与测试', '25'),
        ('    5.1  开发与运行环境', '25'),
        ('    5.2  用户认证模块实现', '25'),
        ('    5.3  个人资料模块实现', '26'),
        ('    5.4  好友管理模块实现', '27'),
        ('    5.5  聊天模块实现', '28'),
        ('    5.6  系统功能测试', '29'),
        ('    5.7  系统性能与安全性测试', '30'),
        ('    5.8  与同类系统的对比', '31'),
        ('结束语', '33'),
        ('致  谢', '34'),
        ('参考文献', '35'),
    ]
    for t, page in toc_items:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(f'{t}')
        set_run_font(r, cn='宋体', en='Times New Roman', size=12)
        r2 = p.add_run('\t' + page)
        set_run_font(r2, cn='宋体', en='Times New Roman', size=12)

    # ========== 节3：正文（第1章起）— 页眉=题目，页码=阿拉伯数字从1重启 ==========
    add_new_section(doc, page_num_fmt='decimal', start=1,
                    header_text=PAPER_TITLE, show_page_number=True)
